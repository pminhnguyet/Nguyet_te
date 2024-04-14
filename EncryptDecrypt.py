import tkinter as tk
from tkinter import filedialog
import docx
from cryptography.fernet import Fernet
import openpyxl

def open_file_dialog():
    text_area_file.config(state=tk.NORMAL)
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        doc = docx.Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        text_area_file.delete('1.0', tk.END) 
        text_area_file.insert(tk.END, content) 
        global data
        data = text_area_file
    text_area_file.config(state=tk.DISABLED)

def update_data():
    global data
    if input_option_var.get() == "1":
        text_area_file.delete('1.0', tk.END) 
        text_area_enter.config(state=tk.NORMAL)
        open_file.config(state=tk.DISABLED)
        text_area_file.config(state=tk.DISABLED)
        data = text_area_enter

    elif input_option_var.get() == "2":
        text_area_enter.delete('1.0', tk.END) 
        text_area_enter.config(state=tk.DISABLED)
        open_file.config(state=tk.NORMAL)
        data = text_area_file

def key_update():
    global key
    global key_1, key_2
    if key_option_var.get() == "1":
        gen_key.config(state=tk.NORMAL)
        gen_key.delete('1.0', tk.END)
        enter_key.config(state=tk.NORMAL)
        gen_key.config(state=tk.DISABLED)
        key_1 = enter_key
        key = key_1.get('1.0', tk.END) 
    elif key_option_var.get() == "2":
        enter_key.delete('1.0', tk.END) 
        enter_key.config(state=tk.DISABLED)
        gen_key.config(state=tk.NORMAL)
        key_2 = Fernet.generate_key()
        gen_key.insert(tk.END, key_2)
        gen_key.config(state=tk.DISABLED)  


def Encrypt():
    if key_option_var.get() == "1":
        key_1 = enter_key
        key = key_1.get('1.0', tk.END) 
    elif key_option_var.get() == "2":
        key = key_2.decode()
    f = Fernet(key)
    enc_show.config(state=tk.DISABLED)   
    global encryptedData
    data_to_encrypt = data.get('1.0', tk.END).encode()
    encryptedData = f.encrypt(data_to_encrypt )
    enc_show.config(state=tk.NORMAL) 
    enc_show.delete('1.0', tk.END)
    enc_show.insert(tk.END, encryptedData.decode())
    enc_show.config(state=tk.DISABLED)
    # print(data_to_encrypt)

def Decrypt():
    #chuan
    if key_option_var.get() == "1":
        key_1 = enter_key
        key = key_1.get('1.0', tk.END) 
    elif key_option_var.get() == "2":
        key = key_2.decode()
    f = Fernet(key)
    dec_show.config(state=tk.DISABLED)   
    global decryptedData
    data_to_decrypt = data.get('1.0', tk.END).encode()
    decryptedData = f.decrypt(data_to_decrypt)
    dec_show.config(state=tk.NORMAL) 
    dec_show.delete('1.0', tk.END)
    dec_show.insert(tk.END, decryptedData.decode())
    dec_show.config(state=tk.DISABLED) 

def export_to_docx():
    data_content = data.get('1.0', tk.END)
    key_content = enter_key.get('1.0', tk.END).strip() if key_option_var.get() == "1" else gen_key.get('1.0', tk.END).strip()
    result_encrypt = enc_show.get('1.0', tk.END).strip() 
    result_decrypt = dec_show.get('1.0', tk.END).strip()

    doc = docx.Document()
    doc.add_heading('Encryption/Decryption Data', level=1)
    doc.add_paragraph(f"Input: {data_content}")
    doc.add_paragraph(f"Key: {key_content}")
    doc.add_paragraph(f"Result of Encrypting: {result_encrypt}")
    doc.add_paragraph(f"Result of Decrypting: {result_decrypt}")

    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if file_path:
        doc.save(file_path)



def export_to_excel():
    data_content = data.get('1.0', tk.END).strip()
    key_content = enter_key.get('1.0', tk.END).strip() if key_option_var.get() == "1" else gen_key.get('1.0', tk.END).strip()
    result_encrypt = enc_show.get('1.0', tk.END).strip() 
    result_decrypt = dec_show.get('1.0', tk.END).strip()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Input", "Key", "Result of Encrypting", "Result of Decrypting"])

    ws.append([data_content, key_content, result_encrypt, result_decrypt ])

    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
    if file_path:
        wb.save(file_path)

def reset():
    text_area_enter.config(state=tk.NORMAL) 
    text_area_file.config(state=tk.NORMAL) 
    enter_key.config(state=tk.NORMAL) 
    gen_key.config(state=tk.NORMAL) 
    enc_show.config(state=tk.NORMAL) 
    dec_show.config(state=tk.NORMAL) 

    text_area_enter.delete('1.0', tk.END)
    text_area_file.delete('1.0', tk.END)
    enter_key.delete('1.0', tk.END)
    gen_key.delete('1.0', tk.END)
    enc_show.delete('1.0', tk.END)
    dec_show.delete('1.0', tk.END)

    # text_area_enter.config(state=tk.NORMAL) 
    text_area_file.config(state=tk.DISABLED) 
    # enter_key.config(state=tk.NORMAL) 
    gen_key.config(state=tk.DISABLED) 
    enc_show.config(state=tk.DISABLED) 
    dec_show.config(state=tk.DISABLED) 

    global data, key_1, key_2
    data = text_area_enter
    key_1 = enter_key
    key_2 = None

    input_option_var.set("1")
    key_option_var.set("1")
    




   

window = tk.Tk()
window.title('Encrypt and Decrypt')
window.geometry("1600x800")

input_choice_label = tk.Label(window, text = "1. Choose method of import data: ")
input_choice_label.place(x = 70, y = 20)

input_option_var = tk.StringVar(window, "1")
input_option1 = tk.Radiobutton(window, text="Enter data for encrypt or decrypt: ", variable = input_option_var, value="1", command=update_data)
input_option1.place(x = 70, y = 50)
text_area_enter = tk.Text(window)
text_area_enter.place(x = 92, y = 80, width = 650, height= 150)
input_option2 = tk.Radiobutton(window, text="Open file: ", variable = input_option_var, value="2", command=update_data)
input_option2.place(x = 770, y = 50)
open_file = tk.Button(window, text = "Open file", command = open_file_dialog, state=tk.DISABLED)
open_file.place( x = 870, y = 50, width = 100)
text_area_file = tk.Text(window,state = tk.DISABLED)
text_area_file.place(x = 792, y = 80, width = 650, height= 150)
data = text_area_enter
# data_bytes = data.encode()



key_choice_label = tk.Label(window, text = "2. Choose key option: ")
key_choice_label.place(x = 70, y = 250)
key_option_var = tk.StringVar(window, "1")
key_option1 = tk.Radiobutton(window, text="Enter key: ", variable = key_option_var, value="1", command=key_update)
key_option1.place(x = 70, y = 270)
enter_key = tk.Text(window)
enter_key.place(x = 92, y = 295, width = 300, height= 35)
key_option2 = tk.Radiobutton(window, text="Generate key: ", variable = key_option_var, value="2", command=key_update)
key_option2.place(x = 422, y = 270)
gen_key = tk.Text(window, state=tk.DISABLED)
gen_key.place(x = 444, y = 295, width = 300, height= 35)
key_1 = enter_key

encOrDec_choice_label = tk.Label(window, text = "3. Choose function: ")
encOrDec_choice_label.place(x = 70, y = 340)

encrypt = tk.Button(window, text = "Encrypt", command = Encrypt)
encrypt.place(x = 92, y = 370, width = 100, height= 40)
enc_show = tk.Text(window, state = tk.DISABLED)
enc_show.place(x = 92, y = 425, width = 650, height= 150)


decrypt = tk.Button(window, text = "Decrypt", command = Decrypt)
decrypt.place(x = 792, y = 370, width = 100, height= 40)
dec_show = tk.Text(window,state = tk.DISABLED)
dec_show.place(x = 792, y = 425, width = 650, height= 150)

export_label_1 = tk.Label(window, text = "Export data to docx file: ")
export_label_1.place(x = 92, y = 600)
export_button_1 = tk.Button(window, text="Export", command=export_to_docx)
export_button_1.place(x = 230, y = 590, width=100, height=40)
export_label_2 = tk.Label(window, text = "Export data to excel file: ")
export_label_2.place(x = 500, y = 600)
export_button_2 = tk.Button(window, text="Export", command=export_to_excel)
export_button_2.place(x = 642, y=590, width=100, height=40)

reset = tk.Button(window, text = "Reset", command= reset)
reset.place(x = 1342, y = 590, width = 100, height= 40)



window.mainloop()